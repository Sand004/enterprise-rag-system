"""DOCX Document Processor"""

import os
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.table import Table
from loguru import logger

from .base import BaseDocumentProcessor


class DOCXProcessor(BaseDocumentProcessor):
    """Process Microsoft Word documents"""
    
    def __init__(self, config: Dict[str, Any]):
        super().__init__(config)
        self.extract_tables = config.get("extract_tables", True)
        self.extract_images = config.get("extract_images", True)
        self.extract_headers_footers = config.get("extract_headers_footers", True)
        self.preserve_formatting = config.get("preserve_formatting", False)
        
    def extract_content(self, file_path: str) -> str:
        """Extract text content from DOCX file"""
        content_parts = []
        
        try:
            # Open the document
            doc = Document(file_path)
            
            # Extract headers if enabled
            if self.extract_headers_footers:
                headers = self._extract_headers(doc)
                if headers:
                    content_parts.append("=== HEADERS ===")
                    content_parts.extend(headers)
                    content_parts.append("")
            
            # Extract main content
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    # Paragraph
                    para = self._extract_paragraph(element, doc)
                    if para.strip():
                        content_parts.append(para)
                        
                elif element.tag.endswith('tbl') and self.extract_tables:
                    # Table
                    table_text = self._extract_table(element, doc)
                    if table_text:
                        content_parts.append("\n" + table_text + "\n")
                        
            # Extract footers if enabled
            if self.extract_headers_footers:
                footers = self._extract_footers(doc)
                if footers:
                    content_parts.append("")
                    content_parts.append("=== FOOTERS ===")
                    content_parts.extend(footers)
                    
            # Extract comments if present
            comments = self._extract_comments(doc)
            if comments:
                content_parts.append("")
                content_parts.append("=== COMMENTS ===")
                content_parts.extend(comments)
                
            # Join all parts
            full_content = "\n".join(content_parts)
            
            # Clean up excessive whitespace
            full_content = re.sub(r'\n{3,}', '\n\n', full_content)
            
            return full_content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {e}")
            raise
            
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX file"""
        metadata = {}
        
        try:
            doc = Document(file_path)
            
            # Core properties
            core_props = doc.core_properties
            
            metadata["title"] = core_props.title or ""
            metadata["author"] = core_props.author or ""
            metadata["subject"] = core_props.subject or ""
            metadata["keywords"] = core_props.keywords or ""
            metadata["category"] = core_props.category or ""
            metadata["comments"] = core_props.comments or ""
            metadata["created"] = str(core_props.created) if core_props.created else ""
            metadata["modified"] = str(core_props.modified) if core_props.modified else ""
            metadata["last_modified_by"] = core_props.last_modified_by or ""
            metadata["revision"] = core_props.revision or 0
            
            # Document statistics
            metadata["page_count"] = self._estimate_page_count(doc)
            metadata["word_count"] = self._count_words(doc)
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
            
            # File information
            metadata["file_size"] = os.path.getsize(file_path)
            metadata["file_name"] = os.path.basename(file_path)
            metadata["file_type"] = "docx"
            
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            metadata["extraction_error"] = str(e)
            
        return metadata
        
    def _extract_paragraph(self, element, doc) -> str:
        """Extract text from a paragraph element"""
        # Get paragraph from element
        for para in doc.paragraphs:
            if para._element == element:
                if self.preserve_formatting:
                    # Preserve basic formatting like bold/italic
                    return self._extract_formatted_text(para)
                else:
                    return para.text
        return ""
        
    def _extract_formatted_text(self, paragraph) -> str:
        """Extract text with basic formatting preserved"""
        formatted_parts = []
        
        for run in paragraph.runs:
            text = run.text
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"_{text}_"
            formatted_parts.append(text)
            
        return "".join(formatted_parts)
        
    def _extract_table(self, element, doc) -> str:
        """Extract text from a table element"""
        # Find the table object
        for table in doc.tables:
            if table._element == element:
                return self._format_table(table)
        return ""
        
    def _format_table(self, table: Table) -> str:
        """Format table data as text"""
        rows = []
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # Get cell text, replacing newlines with spaces
                cell_text = cell.text.replace('\n', ' ').strip()
                cells.append(cell_text)
            rows.append(" | ".join(cells))
            
        # Add table borders
        if rows:
            # Estimate column widths
            header = rows[0]
            separator = "-" * len(header)
            
            formatted_rows = [separator, header, separator]
            formatted_rows.extend(rows[1:])
            formatted_rows.append(separator)
            
            return "\n".join(formatted_rows)
            
        return ""
        
    def _extract_headers(self, doc) -> List[str]:
        """Extract header content from document"""
        headers = []
        
        try:
            for section in doc.sections:
                header = section.header
                if header and header.paragraphs:
                    header_text = "\n".join(p.text.strip() for p in header.paragraphs if p.text.strip())
                    if header_text and header_text not in headers:
                        headers.append(header_text)
        except Exception as e:
            logger.warning(f"Error extracting headers: {e}")
            
        return headers
        
    def _extract_footers(self, doc) -> List[str]:
        """Extract footer content from document"""
        footers = []
        
        try:
            for section in doc.sections:
                footer = section.footer
                if footer and footer.paragraphs:
                    footer_text = "\n".join(p.text.strip() for p in footer.paragraphs if p.text.strip())
                    if footer_text and footer_text not in footers:
                        footers.append(footer_text)
        except Exception as e:
            logger.warning(f"Error extracting footers: {e}")
            
        return footers
        
    def _extract_comments(self, doc) -> List[str]:
        """Extract comments from document"""
        comments = []
        
        # Note: python-docx doesn't have direct comment support
        # This is a placeholder for future implementation
        # Would need to use lxml to parse comments from the XML
        
        return comments
        
    def _estimate_page_count(self, doc) -> int:
        """Estimate page count based on content"""
        # Rough estimation: ~500 words per page
        word_count = self._count_words(doc)
        return max(1, word_count // 500)
        
    def _count_words(self, doc) -> int:
        """Count total words in document"""
        word_count = 0
        
        # Count words in paragraphs
        for para in doc.paragraphs:
            word_count += len(para.text.split())
            
        # Count words in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    word_count += len(cell.text.split())
                    
        return word_count
        
    def extract_images_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract images from DOCX file"""
        images = []
        
        try:
            doc = Document(file_path)
            
            # Access the document's relationships
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_part = rel.target_part
                    image_data = image_part.blob
                    
                    # Get image filename
                    image_name = os.path.basename(image_part.partname)
                    
                    images.append({
                        "name": image_name,
                        "data": image_data,
                        "content_type": image_part.content_type
                    })
                    
        except Exception as e:
            logger.error(f"Error extracting images from DOCX: {e}")
            
        return images
        
    def validate_document(self, file_path: str) -> bool:
        """Validate DOCX document before processing"""
        try:
            # Try to open the document
            doc = Document(file_path)
            
            # Check if document has any content
            has_content = bool(doc.paragraphs or doc.tables)
            
            if not has_content:
                logger.warning(f"Document appears to be empty: {file_path}")
                return False
                
            return True
            
        except Exception as e:
            logger.error(f"Document validation failed: {e}")
            return False