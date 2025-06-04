"""Unit tests for DOCX document processor"""

import os
import tempfile
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

from src.ingestion.docx_processor import DOCXProcessor


class TestDOCXProcessor:
    """Test DOCX processor functionality"""
    
    @pytest.fixture
    def processor(self):
        """Create DOCX processor instance"""
        config = {
            "chunk_size": 1024,
            "chunk_overlap": 256,
            "extract_tables": True,
            "extract_images": True,
            "extract_headers_footers": True,
            "preserve_formatting": False
        }
        return DOCXProcessor(config)
        
    @pytest.fixture
    def sample_docx_file(self):
        """Create a sample DOCX file for testing"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            # Create a simple document
            doc = Document()
            
            # Add title
            doc.add_heading("Test Document", 0)
            
            # Add paragraphs
            doc.add_paragraph("This is the first paragraph of the test document.")
            doc.add_paragraph("This is the second paragraph with some content.")
            
            # Add a table
            table = doc.add_table(rows=3, cols=3)
            table.cell(0, 0).text = "Header 1"
            table.cell(0, 1).text = "Header 2"
            table.cell(0, 2).text = "Header 3"
            table.cell(1, 0).text = "Data 1"
            table.cell(1, 1).text = "Data 2"
            table.cell(1, 2).text = "Data 3"
            
            # Add another paragraph
            doc.add_paragraph("This is the final paragraph after the table.")
            
            # Set some metadata
            doc.core_properties.title = "Test Document"
            doc.core_properties.author = "Test Author"
            doc.core_properties.subject = "Testing"
            doc.core_properties.keywords = "test, docx, processor"
            
            # Save the document
            doc.save(f.name)
            temp_path = f.name
            
        yield temp_path
        
        # Cleanup
        os.unlink(temp_path)
        
    def test_docx_processor_initialization(self, processor):
        """Test DOCX processor initialization"""
        assert processor.chunk_size == 1024
        assert processor.chunk_overlap == 256
        assert processor.extract_tables is True
        assert processor.extract_images is True
        assert processor.extract_headers_footers is True
        assert processor.preserve_formatting is False
        
    def test_extract_content(self, processor, sample_docx_file):
        """Test content extraction from DOCX"""
        content = processor.extract_content(sample_docx_file)
        
        assert "Test Document" in content
        assert "This is the first paragraph" in content
        assert "This is the second paragraph" in content
        assert "Header 1" in content
        assert "Data 1" in content
        assert "This is the final paragraph" in content
        
    def test_extract_metadata(self, processor, sample_docx_file):
        """Test metadata extraction from DOCX"""
        metadata = processor.extract_metadata(sample_docx_file)
        
        assert metadata["title"] == "Test Document"
        assert metadata["author"] == "Test Author"
        assert metadata["subject"] == "Testing"
        assert metadata["keywords"] == "test, docx, processor"
        assert metadata["file_type"] == "docx"
        assert metadata["file_name"].endswith(".docx")
        assert metadata["paragraph_count"] > 0
        assert metadata["table_count"] > 0
        assert metadata["word_count"] > 0
        
    def test_table_extraction(self, processor, sample_docx_file):
        """Test table extraction functionality"""
        content = processor.extract_content(sample_docx_file)
        
        # Check that table content is present
        assert "Header 1 | Header 2 | Header 3" in content
        assert "Data 1 | Data 2 | Data 3" in content
        
    def test_word_count(self, processor):
        """Test word counting functionality"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.add_paragraph("This is a test with exactly ten words in it.")
            doc.save(f.name)
            
            word_count = processor._count_words(doc)
            assert word_count == 10
            
            os.unlink(f.name)
            
    def test_empty_document(self, processor):
        """Test handling of empty documents"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            doc.save(f.name)
            
            # Should not validate
            is_valid = processor.validate_document(f.name)
            assert is_valid is False
            
            os.unlink(f.name)
            
    def test_document_with_formatting(self):
        """Test document with formatting preservation"""
        config = {
            "chunk_size": 1024,
            "chunk_overlap": 256,
            "preserve_formatting": True
        }
        processor = DOCXProcessor(config)
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc = Document()
            para = doc.add_paragraph()
            para.add_run("Normal text ")
            para.add_run("bold text").bold = True
            para.add_run(" and ")
            para.add_run("italic text").italic = True
            doc.save(f.name)
            
            content = processor.extract_content(f.name)
            
            assert "Normal text" in content
            assert "**bold text**" in content
            assert "*italic text*" in content
            
            os.unlink(f.name)
            
    def test_process_document_pipeline(self, processor, sample_docx_file):
        """Test complete document processing pipeline"""
        processed_doc = processor.process_document(
            sample_docx_file,
            user_context={"user_id": "test_user"}
        )
        
        assert processed_doc.id
        assert processed_doc.source_path == sample_docx_file
        assert len(processed_doc.content) > 0
        assert len(processed_doc.chunks) > 0
        assert processed_doc.metadata["file_type"] == "docx"
        assert processed_doc.checksum
        
        # Check chunks
        for chunk in processed_doc.chunks:
            assert chunk.document_id == processed_doc.id
            assert len(chunk.content) > 0
            assert chunk.metadata["user_context"]["user_id"] == "test_user"
            
    @patch('src.ingestion.docx_processor.Document')
    def test_extraction_error_handling(self, mock_doc_class, processor):
        """Test error handling during extraction"""
        # Mock Document to raise an exception
        mock_doc_class.side_effect = Exception("Test error")
        
        with pytest.raises(Exception) as exc_info:
            processor.extract_content("fake_path.docx")
            
        assert "Test error" in str(exc_info.value)